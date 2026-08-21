"""Test cho luồng tài liệu docx đính kèm chat: trích text, render khối prompt, kiểm URL.

Chỉ test các hàm thuần (không DB, không mạng). Phần join vào lịch sử pha answer và endpoint
được kiểm bằng tay/hệ tích hợp vì cần DB thật.
"""

from datetime import datetime, timedelta, timezone
from io import BytesIO

import pytest
from docx import Document

from apps.chat.utils.attachment import (
    ERR_DOC_PARSE_FAILED,
    AttachmentContent,
    dedupe_file_urls,
    extract_docx_text,
    render_attachment_block,
    render_attachment_blocks,
    split_char_budget,
)
from apps.datasource.utils.excel_import import ExcelImportError
from apps.datasource.utils.remote_file import validate_source_url
from common.core.config import settings


def _make_docx(paragraphs=(), table_rows=None) -> bytes:
    """Dựng một file .docx trong RAM: các đoạn văn trước, một bảng (nếu có) sau."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                table.cell(r, c).text = cell
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractDocxText:
    def test_paragraphs_and_table_markdown(self):
        data = _make_docx(
            paragraphs=['Đoạn mở đầu.', 'Đoạn thứ hai.'],
            table_rows=[['Tên', 'Giá trị'], ['Thu ngân sách', '1.234']],
        )
        text, truncated = extract_docx_text(data, max_chars=10_000)
        assert not truncated
        assert 'Đoạn mở đầu.' in text
        # Bảng thành markdown table: header + dòng phân cách + dòng dữ liệu
        assert '| Tên | Giá trị |' in text
        assert '| --- | --- |' in text
        assert '| Thu ngân sách | 1.234 |' in text
        # Đoạn văn đứng trước bảng theo đúng thứ tự tài liệu
        assert text.index('Đoạn thứ hai.') < text.index('| Tên |')

    def test_pipe_in_cell_is_escaped(self):
        data = _make_docx(table_rows=[['a|b', 'c'], ['d', 'e']])
        text, _ = extract_docx_text(data, max_chars=10_000)
        assert 'a\\|b' in text

    def test_truncation_stops_at_limit(self):
        data = _make_docx(paragraphs=['x' * 50] * 10)
        text, truncated = extract_docx_text(data, max_chars=120)
        assert truncated
        assert len(text) <= 120
        # Cắt tại ranh giới đoạn: không có đoạn nào bị đứt giữa chừng
        assert all(len(line) == 50 for line in text.split('\n') if line)

    def test_empty_paragraphs_skipped(self):
        data = _make_docx(paragraphs=['Một.', '', '   ', 'Hai.'])
        text, _ = extract_docx_text(data, max_chars=10_000)
        assert text.split('\n') == ['Một.', 'Hai.']

    def test_not_a_docx_raises_parse_failed(self):
        with pytest.raises(ExcelImportError) as exc_info:
            extract_docx_text(b'this is not a zip file', max_chars=1000)
        assert exc_info.value.error_code == ERR_DOC_PARSE_FAILED


class TestRenderAttachmentBlock:
    def test_basic_block(self):
        block = render_attachment_block('bao_cao.docx', 'Nội dung.', False, max_chars=1000)
        assert block.startswith('<attached-document filename="bao_cao.docx">')
        assert block.endswith('</attached-document>')
        assert 'Nội dung.' in block
        # Không cắt gì thì không có lời báo cắt
        assert 'KHÔNG được gửi' not in block

    def test_cut_at_render_has_notice(self):
        block = render_attachment_block('a.docx', 'x' * 100, False, max_chars=10)
        assert 'x' * 10 in block
        assert 'x' * 11 not in block
        # Quy tắc repo: cắt là phải nói cho LLM biết đã cắt
        assert 'KHÔNG được gửi cho bạn' in block

    def test_source_truncated_has_notice(self):
        block = render_attachment_block('a.docx', 'ngắn', True, max_chars=1000)
        assert 'cắt bớt từ lúc trích xuất' in block

    def test_filename_sanitized(self):
        block = render_attachment_block('b"o<c>a.docx', 'x', False, max_chars=100)
        assert 'filename="b_o_c_a.docx"' in block

    def test_indent_applied(self):
        block = render_attachment_block('a.docx', 'dòng', False, max_chars=100, indent='  ')
        for line in block.split('\n'):
            assert line.startswith('  ')


class TestSplitCharBudget:
    """Ngân sách ký tự của một lượt là trần TỔNG — chia thế nào cũng không được vượt."""

    def test_chia_deu_khi_moi_file_deu_dai(self):
        assert split_char_budget([1000, 1000], 100) == [50, 50]

    def test_file_ngan_khong_om_suat_thua(self):
        # File 10 ký tự chỉ lấy 10, 90 còn lại chia cho hai file đói thay vì bốc hơi
        quotas = split_char_budget([10, 1000, 1000], 100)
        assert quotas == [10, 45, 45]
        assert sum(quotas) <= 100

    def test_tat_ca_deu_vua_thi_khong_ai_bi_cat(self):
        assert split_char_budget([10, 20, 30], 100) == [10, 20, 30]

    def test_khong_bao_gio_vuot_tran_tong(self):
        for total in (0, 1, 7, 99, 100_000):
            for lengths in ([5], [5, 5], [1, 2, 3, 4, 5], [10_000] * 5):
                assert sum(split_char_budget(lengths, total)) <= total

    def test_danh_sach_rong(self):
        assert split_char_budget([], 100) == []


class TestRenderAttachmentBlocks:
    def _att(self, name, content, truncated=False):
        return AttachmentContent(filename=name, content=content, truncated=truncated)

    def test_giu_thu_tu_client_gui(self):
        blocks = render_attachment_blocks(
            [self._att('a.docx', 'AAA'), self._att('b.docx', 'BBB')], total_max_chars=1000)
        assert blocks.index('a.docx') < blocks.index('b.docx')
        assert blocks.count('<attached-document') == 2

    def test_ngan_sach_la_tran_tong_chu_khong_phai_moi_file(self):
        # Hai file, mỗi file 100 ký tự, trần tổng 100 -> mỗi file chỉ được 50
        blocks = render_attachment_blocks(
            [self._att('a.docx', 'x' * 100), self._att('b.docx', 'y' * 100)], total_max_chars=100)
        assert 'x' * 50 in blocks and 'x' * 51 not in blocks
        assert 'y' * 50 in blocks and 'y' * 51 not in blocks
        # Cắt thì phải báo, cả hai khối đều bị cắt nên có hai lời báo
        assert blocks.count('KHÔNG được gửi cho bạn') == 2

    def test_danh_sach_rong_tra_chuoi_rong(self):
        assert render_attachment_blocks([], total_max_chars=1000) == ''

    def test_indent_ap_cho_moi_khoi(self):
        blocks = render_attachment_blocks(
            [self._att('a.docx', 'p'), self._att('b.docx', 'q')], total_max_chars=1000, indent='  ')
        for line in blocks.split('\n'):
            assert line.startswith('  ')


class TestDedupeFileUrls:
    def test_bo_trung_giu_lan_xuat_hien_dau(self):
        assert dedupe_file_urls(['a', 'b', 'a']) == [(0, 'a'), (1, 'b')]

    def test_vi_tri_la_vi_tri_trong_mang_goc(self):
        # Sau khi bỏ trùng, 'c' vẫn phải mang số 3 — đó là số client dùng để tìm file hỏng
        assert dedupe_file_urls(['a', 'a', 'a', 'c']) == [(0, 'a'), (3, 'c')]

    def test_khoang_trang_thua_van_tinh_la_trung(self):
        assert dedupe_file_urls(['http://x/a.docx', ' http://x/a.docx ']) == [(0, 'http://x/a.docx')]

    def test_url_khac_nhau_thi_giu_ca_hai(self):
        assert len(dedupe_file_urls(['http://x/a.docx', 'http://x/b.docx'])) == 2


class TestValidateSourceUrlForDocx:
    HOST = 'minio.local:9000'

    def _url(self, name='van_ban.docx', expires_in: int | None = None) -> str:
        url = f'http://{self.HOST}/bucket/{name}?X-Amz-Signature=abc'
        if expires_in is not None:
            signed_at = datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')
            url += f'&X-Amz-Date={signed_at}&X-Amz-Expires={expires_in}'
        return url

    @pytest.fixture(autouse=True)
    def _allow_host(self, monkeypatch):
        monkeypatch.setattr(settings, 'FILE_DOWNLOAD_ALLOWED_HOSTS', self.HOST)

    def test_docx_extension_accepted(self):
        src = validate_source_url(self._url(), allowed_extensions=('docx',))
        assert src.extension == 'docx'
        assert src.object_key == '/bucket/van_ban.docx'

    def test_default_extensions_reject_docx(self):
        # Hành vi cũ của luồng excel không đổi: đuôi docx vẫn bị chặn khi dùng mặc định
        with pytest.raises(ExcelImportError) as exc_info:
            validate_source_url(self._url())
        assert exc_info.value.error_code == 'URL_BAD_EXTENSION'

    def test_short_ttl_ok_with_low_min_ttl(self):
        # URL còn sống 60s: luồng chat (min_ttl thấp) nhận, luồng excel (mặc định 300s) từ chối
        url = self._url(expires_in=60)
        src = validate_source_url(url, allowed_extensions=('docx',), min_ttl_seconds=10)
        assert src.expires_at is not None
        with pytest.raises(ExcelImportError) as exc_info:
            validate_source_url(url, allowed_extensions=('docx', 'xlsx', 'xls', 'csv'))
        assert exc_info.value.error_code == 'URL_EXPIRED'
